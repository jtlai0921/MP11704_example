import datetime
import tkinter
import tkinter.scrolledtext
import tkinter.messagebox
import tkinter.filedialog
import tkinter.ttk
import tkinter.simpledialog
import socket
import sqlite3
import random
import threading
import time
import struct
import os
import sys
import string
import re

## =====================安裝與升級擴展庫功能程式碼開始==============================
#先把pip升級到最新版本
path = '"'+os.path.dirname(sys.executable)+'\\scripts\\pip" install --upgrade pip'
os.system(path)
#匯入必要的擴展庫
#docx擴展庫
try:
    import docx
except:
    path = '"'+os.path.dirname(sys.executable)+'\\scripts\\pip" install python-docx'
    os.system(path)
    import docx
#xlrd擴展庫
try:
    import xlrd
except:
    path = '"'+os.path.dirname(sys.executable)+'\\scripts\\pip" install xlrd'
    os.system(path)
    import xlrd
#openpyxl擴展庫
try:
    import openpyxl
except:
    path = '"'+os.path.dirname(sys.executable)+'\\scripts\\pip" install openpyxl'
    os.system(path)
    import openpyxl
#pillow擴展庫
try:
    from PIL import ImageGrab
except:
    path = '"'+os.path.dirname(sys.executable)+'\\scripts\\pip" install pillow'
    os.system(path)
    from PIL import ImageGrab
## =====================安裝與升級擴展庫功能程式碼結束==============================
    
root = tkinter.Tk()
#root.config(width=360)
#root.config(height=260)
root.geometry('360x420+400+200')
# 不允許改變視窗大小
root.resizable(False, False)
root.title('課堂教學管理系統v2.0---董付國')

# 抓取本機IP
serverIP = socket.gethostbyname(socket.gethostname())
if serverIP.startswith('127.0.'):
    addrs = socket.getaddrinfo(socket.gethostname(  ),None,0,socket.SOCK_STREAM)
    addrs = [x[4][0] for x in addrs]
    serverIP = [x for x in addrs if ':' not in x][0]

# 廣播，傳送伺服端IP位址
def sendServerIP():
    # 建立socket物件
    sock = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    sock.setsockopt(socket.SOL_SOCKET, socket.SO_BROADCAST, 1)
    while True:        
        # 255表示廣播位址
        IP = serverIP[:serverIP.rindex('.')]+'.255'
        # 3秒鐘傳送一次訊息
        sock.sendto('ServerIP'.encode(), (IP, 5000))
        time.sleep(3)
thread_sendServerIP = threading.Thread(target=sendServerIP)
thread_sendServerIP.start()

# 關閉程式時，取消點名、收作業、接收提問以及客戶端查詢等狀態，避免一直佔用連接埠
def closeWindow():
    # 結束點名
    if int_canDianming.get() == 1:
        int_canDianming.set(0)
        
    # 結束收作業
    if int_zuoye.get() == 1:
        int_zuoye.set(0)
        
    # 結束學生主動提問
    if int_xueshengTiwen.get() == 1:
        int_xueshengTiwen.set(0)
        
    # 結束螢幕廣播
    global broadcasting
    broadcasting = False
    try:
        sockBroadCast.close()
    except:
        pass
##    #結束服務狀態
##    if int_server.get() == 1:
##        int_server.set(0)
    root.destroy()
root.protocol('WM_DELETE_WINDOW', closeWindow)


## =====================通用功能程式碼開始==============================
class Common:
    # 查詢資料庫，抓取學生專業清單
    def getZhuanye():
        with sqlite3.connect('database.db') as conn:
            cur = conn.cursor()
            cur.execute('select distinct(zhuanye) from students')
            temp = cur.fetchall()
        xueshengZhuanye = []
        for line in temp:
            xueshengZhuanye.append(line[0])
        return xueshengZhuanye

    # 抓取指定專業的學生名單
    def getXuehaoXingming(zhuanye):
        with sqlite3.connect('database.db') as conn:
            cur = conn.cursor()
            cur.execute("select xuehao,xingming from students where zhuanye='"+zhuanye+"' order by xuehao")
            temp = cur.fetchall()
        xueshengXinxi = []
        for line in temp:
            xueshengXinxi.append(line[0]+','+line[1])
        return xueshengXinxi

    # 抓取指定學號的出勤次數
    def getChuqinCishu(xuehao):
        with sqlite3.connect('database.db') as conn:
            cur = conn.cursor()
            cur.execute("select count(xuehao) from dianming where xuehao='"+xuehao+"'")
            temp = cur.fetchall()
        return temp[0][0]

    # 抓取指定學號的學生提問總得分
    def getTiwenDefen(xuehao):
        with sqlite3.connect('database.db') as conn:
            cur = conn.cursor()
            cur.execute("select sum(defen) from tiwen where xuehao='"+xuehao+"'")
            temp = cur.fetchall()
        return temp[0][0]
    
    # 抓取指定學號的學生主動提問次數
    def getZhudongTiwenCishu(xuehao):
        with sqlite3.connect('database.db') as conn:
            cur = conn.cursor()
            cur.execute("select count(xuehao) from xueshengtiwen where xuehao='"+xuehao+"' and wenti not like '老師回覆%'")
            temp = cur.fetchall()
        return temp[0][0]

    # 查看學生線上考試得分
    def getKaoshiDefen(xuehao):
        with sqlite3.connect('database.db') as conn:
            cur = conn.cursor()
            cur.execute("select count(xuehao) from kaoshi where xuehao='"+xuehao+"' and shifouzhengque='Y'")
            temp = cur.fetchall()
        return temp[0][0]

    # 抓取指定SQL語句查詢結果
    def getDataBySQL(sql):
        with sqlite3.connect('database.db') as conn:
            cur = conn.cursor()
            cur.execute(sql)
            result = cur.fetchall()
        return result
    
    # 執行SQL語句
    def doSQL(sql):
        with sqlite3.connect('database.db') as conn:
            cur = conn.cursor()
            cur.execute(sql)
            conn.commit()

    # 目前日期時間，格式為"年-月-日 時:分:秒"
    def getCurrentDateTime():
        return str(datetime.datetime.now())[:19]

    # 前一個半小時的時間，主要用來避色重複點名
    def getStartDateTime():
        now = datetime.datetime.now()
        now = now + datetime.timedelta(minutes=-90)
        return str(now)[:19]
## =====================通用功能程式碼結束==============================

    
# 控制和檢查本軟體是否已註冊，本軟體免費試用，不需要註冊
int_zhuce = tkinter.IntVar(root, value=0)
def isZhuce():
    int_zhuce.set(1)
    
isZhuce()

## =====================匯入學生資料功能程式碼開始==============================
def buttonImportXueshengXinxiClick():
    # 如果還沒有註冊，拒絕執行
    if int_zhuce.get() == 0:
        tkinter.messagebox.showerror('很抱歉', '請聯繫作者進行軟體註冊！')
        return
    filename = tkinter.filedialog.askopenfilename(title='請選擇Excel檔案',
                                                  filetypes=[('Excel Files','*.xls')])
    if filename:
        # 讀取資斛並匯入資料庫
        workbook = xlrd.open_workbook(filename=filename)
        sheet1 = workbook.sheet_by_index(0)
        # Excel檔案必須包含4欄，分別是學號、姓名、專業年級、課程名稱
        if sheet1.ncols != 4:
            tkinter.messagebox.showerror(title='很抱歉', message='Excel檔案格式不對！')
            return

        # 巡訪Excel檔案每一列
        for rowIndex in range(1, sheet1.nrows):
            row = sheet1.row(rowIndex)
            sql = "insert into students(xuehao,xingming,zhuanye,kecheng) values('"\
                  + "','".join(map(lambda item:str(item.value).strip(), row)) + "')"
##            sql = "insert into students(xuehao,xingming,zhuanye,kecheng) values('"\
##                  +str(row[0].value).strip()+"','"+str(row[1].value)+"','"\
##                  +str(row[2].value)+"','"+str(row[3].value)+"')"
            Common.doSQL(sql)
        tkinter.messagebox.showinfo(title='恭喜', message='匯入成功')
buttonImportXueshengXinxi = tkinter.Button(root, text='匯入學生資料', command=buttonImportXueshengXinxiClick)
buttonImportXueshengXinxi.place(x=20, y=20, height=30, width=100)
## =====================匯入學生資料功能程式碼結束==============================


## =====================查詢學生資料功能程式碼開始==============================
# 查詢學生資訊，可以給認真聽課的同學加分
class windowChakanXueshengXinxi:
    def __init__(self, root, myTitle):
        self.top = tkinter.Toplevel(root, width=350, height=400)
        self.top.title(myTitle)
        self.top.attributes('-topmost', 1)

        #以下拉式清單顯非學生專業，呼叫通用功能類別的方法抓取其專業
        xueshengZhuanye = Common.getZhuanye()
        comboboxZhuanye = tkinter.ttk.Combobox(self.top, values=xueshengZhuanye)
        comboboxZhuanye.place(x=20, y=20, height=20, width=100)
        def buttonChakanClick():
            zhuanye = comboboxZhuanye.get()
            if not zhuanye:
                tkinter.messagebox.showerror(title='很抱歉', message='請選擇一個專業！')
                return
            
            # 根據選擇的專業，抓取該專業所有學生名單，格式為"學號,姓名"
            temp = Common.getXuehaoXingming(zhuanye)
            # 刪除原有的記錄
            for row in treeXueshengMingdan.get_children():
                treeXueshengMingdan.delete(row)

            for iii, student in enumerate(temp):
                #切割學號和姓名
                treeXueshengMingdan.insert('', iii, values=student.split(','))
        buttonChakan = tkinter.Button(self.top, text='查詢', command=buttonChakanClick)
        buttonChakan.place(x=130, y=20, height=20, width=40)
        
        # 建立表格，設定表頭，show="headings"用來隱藏樹狀元件的預設首行
        self.frame = tkinter.Frame(self.top)
        self.frame.place(x=20, y=50, width=200, height=280)
        # 垂直捲軸
        scrollBar = tkinter.Scrollbar(self.frame)
        scrollBar.pack(side=tkinter.RIGHT, fill=tkinter.Y)
        # 以樹狀元件實作表格
        treeXueshengMingdan = tkinter.ttk.Treeview(self.frame,
                                                   columns=('col1', 'col2'),
                                                   show="headings",
                                                   yscrollcommand = scrollBar.set)
        treeXueshengMingdan.column('col1', width=90, anchor='center')
        treeXueshengMingdan.column('col2', width=90, anchor='center')
        treeXueshengMingdan.heading('col1', text='學號')
        treeXueshengMingdan.heading('col2', text='姓名')
        
        # 擊擊某筆學生資料，雙擊實現離線點名
        def onDBClick(event):
            selected = treeXueshengMingdan.selection()
            if not selected:
                tkinter.messagebox.showerror('很抱歉', '請挑選學生！')
                return
            
            item = selected[0]
            xuehaoDianming, xingmingDianming = treeXueshengMingdan.item(item, 'values')
            
            currentTime = Common.getCurrentDateTime()
            # 抓取一個半小時之前的時間
            startTime = Common.getStartDateTime()
            # 查詢是否已經點過名，避免一個半小時內重複點名
            sqlShifouChongfuDianming = "select count(xuehao) from dianming where xuehao='"\
                                       +xuehaoDianming+"' and shijian >='"+startTime+"'"
            if Common.getDataBySQL(sqlShifouChongfuDianming)[0][0] != 0:
                tkinter.messagebox.showerror('很抱歉', xuehaoDianming+','+xingmingDianming+'重複點名！')
                return
            # 點名
            sqlDianming = "insert into dianming(xuehao,shijian) values('"+xuehaoDianming+"','"+currentTime+"')"
            Common.doSQL(sqlDianming)
            tkinter.messagebox.showinfo('恭喜', xuehaoDianming+','+xingmingDianming+'  點名成功')
        treeXueshengMingdan.bind("<Double-1>", onDBClick)
        treeXueshengMingdan.pack(side=tkinter.LEFT, fill=tkinter.Y)
        # 結合樹狀元件與重直捲軸
        scrollBar.config(command=treeXueshengMingdan.yview)
        
        def scoreAddSub(score):
            # 抓取目前選擇項目
            selected = treeXueshengMingdan.selection()
            if not selected:
                tkinter.messagebox.showerror('很抱歉', '請選擇學生！')
                return
            
            item = selected[0]
            xuehaoJiafen = treeXueshengMingdan.item(item, 'values')[0]
            sqlJiafen = "insert into tiwen(xuehao,shijian,defen) values('"\
                        +xuehaoJiafen+"','" + Common.getCurrentDateTime()+"'," + score +")"
            Common.doSQL(sqlJiafen)
            
        def buttonJiafenClick():
            # 首先挑選一個學生，然後加分
            # 為該名學生增加一個提問得分，5分
            scoreAddSub('5')
            tkinter.messagebox.showinfo('恭喜', '加分成功！')
        buttonJiafen = tkinter.Button(self.top, text='聽課認真加分', command=buttonJiafenClick)
        buttonJiafen.place(x=30, y=350, height=20, width=100)
        
        def buttonJianfenClick():
            # 首先選擇一個學生，然後減分
            # 為該名學生增加一個提問得分，-5分
            scoreAddSub('-5')
            tkinter.messagebox.showinfo('恭喜', '減分成功')
        buttonJianfen = tkinter.Button(self.top, text='聽課不認真減分', command=buttonJianfenClick)
        buttonJianfen.place(x=140, y=350, height=20, width=100)
        labelTishi = tkinter.Label(self.top, text='溫馨提示：雙擊表格中某位學生可以離線點名，或者補點名。', fg='red')
        labelTishi.place(x=10, y=380, height=20)
        
def buttonChakanXueshengxinxiClick():
    # 如果還沒有註冊，拒絕執行
    if int_zhuce.get() == 0:
        tkinter.messagebox.showerror('很抱歉', '請聯繫作者進行軟體註冊！')
        return
    buttonChakanXueshengxinxi['state'] = 'disabled'
    w1 = windowChakanXueshengXinxi(root, '查詢學生資料')
    buttonChakanXueshengxinxi.wait_window(w1.top)
    buttonChakanXueshengxinxi['state'] = 'normal'
buttonChakanXueshengxinxi = tkinter.Button(root, text='查詢學生資料', command=buttonChakanXueshengxinxiClick)
buttonChakanXueshengxinxi.place(x=130, y=20, height=30, width=100)
## =====================查詢學生資料功能程式碼結束==============================


def buttonIPClick():
    #抓取並輸出本機IP位址
    #如果還沒有註冊，拒絕執行
    if int_zhuce.get() == 0:
        tkinter.messagebox.showerror('很抱歉', '請聯繫作者進行軟體註冊！')
        return
    tkinter.messagebox.showinfo(title='本機IP位址', message=serverIP)
buttonIP = tkinter.Button(root, text='查詢本機IP位址', command=buttonIPClick)
buttonIP.place(x=240, y=20, height=30, width=100)

## =====================線上點名功能程式碼開始==============================
# 控制是否可以點名，1表示可以，0表示不可以
int_canDianming = tkinter.IntVar(root, value=0)
def thread_Dianming():
    # 開始監聽
    global sockDianming
    sockDianming = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    sockDianming.bind(('', 30300))
    sockDianming.listen(200)
    while int_canDianming.get()==1:
        try:
            #接受客戶端連接
            conn, addr = sockDianming.accept()
        except:
            continue
        data = conn.recv(1024).decode()
        try:
            #客戶端送來的訊息格式為：學號,姓名,MAC位址
            xuehao, xingming, mac = data.split(',')
        except:
            conn.sendall('notmatch'.encode())
            conn.close()
            continue
        # 防SQL注入
        xuehao = re.sub(r'[;"\'=]', '', xuehao)
        xingming = re.sub(r'[;"\'=]', '', xingming)
        # 首先檢查學號與姓名是否相符，並且與資料庫的學生資料一致
        sqlIfMatch = "select count(xuehao) from students where xuehao='" + xuehao + "' and xingming='" + xingming + "'"
        if Common.getDataBySQL(sqlIfMatch)[0][0] != 1:
            conn.sendall('notmatch'.encode())
            conn.close()
        else:
            # 記錄該學生點名資料：學號、姓名、時間，並回應至客戶端點名成功，然後客戶端關閉連接
            currentTime = Common.getCurrentDateTime()
            # 抓取一個半小時之前的時間
            startTime = Common.getStartDateTime()
            # 查詢是否已經點過名，避免一個半小時內重複點名
            sqlShifouChongfuDianming = "select count(xuehao) from dianming where xuehao='"\
                                       +xuehao+"' and shijian >='"+startTime+"'"
                        
            if Common.getDataBySQL(sqlShifouChongfuDianming)[0][0] != 0:
                conn.sendall('repeat'.encode())
                conn.close()
            else:
                #檢查是否乏替點名，根據學生端IP位址識別
                sqlShifouDaiDianming = "select count(ip) from dianming where ip='"\
                                       +addr[0]+"' and shijian >='"+startTime+"'"
                sqlMacShifouChongfu = "select count(mac) from dianming where mac='"\
                                      +mac+"' and shijian>='"+startTime+"'"
                
                if Common.getDataBySQL(sqlShifouDaiDianming)[0][0] != 0 \
                   or Common.getDataBySQL(sqlMacShifouChongfu)[0][0] != 0:
                    conn.sendall('daidianming'.encode())
                    conn.close()
                else:
                    #點名
                    sqlDianming = "insert into dianming(xuehao,shijian,ip,mac) values('"\
                                  +xuehao+"','"+currentTime+"','"+addr[0]+"','"+mac+"')"
                    Common.doSQL(sqlDianming)
                    conn.sendall('ok'.encode())
                    conn.close()
    sockDianming.close()
    sockDianming = None

# 開始點名
def buttonStartDianmingClick():
    # 如果還沒有註冊，拒絕執行
    if int_zhuce.get() == 0:
        tkinter.messagebox.showerror('很抱歉', '請聯繫作者進行軟體註冊！')
        return
    if int_zuoye.get() == 1:
        tkinter.messagebox.showerror('很抱歉', '現在正在收作業！')
        return
    if int_canDianming.get() == 1:
        tkinter.messagebox.showerror('很抱歉', '現在正在點名！')
        return
    tkinter.messagebox.showinfo('恭喜', '設定成功，現在開始點名！')
    #開始點名
    int_canDianming.set(1)
    global tDianming_id
    t = threading.Thread(target=thread_Dianming)
    t.start()
    tDianming_id = t.ident
buttonStartDianming = tkinter.Button(root, text='開始點名', command=buttonStartDianmingClick)
buttonStartDianming.place(x=20, y=60, height=30, width=100)

def buttonStopDianmingClick():
    # 如果還沒有註冊，拒絕執行
    if int_zhuce.get() == 0:
        tkinter.messagebox.showerror('很抱歉', '請聯繫作者進行軟體註冊！')
        return
    if int_canDianming.get() == 0:
        tkinter.messagebox.showerror('很抱歉', '還沒開始點名！')
        return

    #停止點名
    int_canDianming.set(0)
    sockDianming.close()
    time.sleep(0.1)
    sql = 'select zhuanye from students where xuehao=(select xuehao from dianming where shijian<="'\
          + Common.getCurrentDateTime() + '"  order by shijian desc limit 1)'
    currentZhuanye = Common.getDataBySQL(sql)[0][0]
    sql = 'select count(zhuanye) from students where zhuanye="' + currentZhuanye + '"'
    totalRenshu = Common.getDataBySQL(sql)[0][0]

    sql = 'select count(xuehao) from dianming where shijian<="'+Common.getCurrentDateTime()\
          +'" and shijian>="' + Common.getStartDateTime() + '"'
    totalShidao = Common.getDataBySQL(sql)[0][0]
    
    message = '設定成功，現在停止點名!\n目前點名課程：'+currentZhuanye\
              +'\n應到人數：'+str(totalRenshu) + '\n實到人數：' + str(totalShidao)
    tkinter.messagebox.showinfo('恭喜', message)
buttonStopDianming = tkinter.Button(root, text='結束點名', command=buttonStopDianmingClick)
buttonStopDianming.place(x=130, y=60, height=30, width=100)
## =====================線上點名功能程式碼結束==============================


## =====================查詢學生出勤情況功能程式碼開始==============================
int_windowChakanKaoqinXinxi = tkinter.IntVar(root, value=0)
class windowChakanKaoqinXinxi:
    def __init__(self, root, myTitle):
        self.top = tkinter.Toplevel(root, width=320, height=400)
        self.top.title(myTitle)
        self.top.attributes('-topmost', 1)
        # 下拉式清單選擇專業，或文字框查詢特定學生的出勤情況
        xueshengZhuanye = Common.getZhuanye()
        comboboxZhuanye = tkinter.ttk.Combobox(self.top, values=xueshengZhuanye)
        def comboboxZhuanyeChange(event):
            zhuanye = comboboxZhuanye.get()
            if zhuanye:
                xueshengs = Common.getXuehaoXingming(zhuanye)
                comboboxXuehao['values'] = xueshengs                
        comboboxZhuanye.bind('<<ComboboxSelected>>', comboboxZhuanyeChange)
        comboboxZhuanye.place(x=20, y=20, height=20, width=130)

        # 輸出該課程所有學生的出勤情況
        def chakanZhuanye():
            zhuanye = comboboxZhuanye.get()
            if not zhuanye:
                tkinter.messagebox.showerror('很抱歉', '請選擇一個專業！')
                return
            sql = "select students.xuehao,students.xingming,dianming.shijian from students,dianming where students.xuehao=dianming.xuehao and students.zhuanye='"+zhuanye+"' order by students.xuehao"
            temp = Common.getDataBySQL(sql)

            # 刪除原有的所有列
            for row in treeXueshengMingdan.get_children():
                treeXueshengMingdan.delete(row)

            # 插入新資料
            for iii, student in enumerate(temp):
                treeXueshengMingdan.insert('', iii, values=(student[0], student[1], student[2]))
        buttonZhuanye = tkinter.Button(self.top, text='依專業查詢', command=chakanZhuanye)
        buttonZhuanye.place(x=160, y=20, height=20, width=80)
        comboboxXuehao = tkinter.ttk.Combobox(self.top,)
        comboboxXuehao.place(x=20, y=60, height=20, width=130)

        # 依學號查詢學生出勤情況
        def chakanXuehao():
            # 刪除原有的所有列
            for row in treeXueshengMingdan.get_children():
                treeXueshengMingdan.delete(row)
            xueshengXinxi = comboboxXuehao.get()
            if not xueshengXinxi:
                tkinter.messagebox.showerror('很抱歉', '請挑選學生！')
                return
            xuehaoKaoqin = xueshengXinxi.split(',')[0]
            sql = "select students.xuehao,students.xingming,dianming.shijian from students,dianming where students.xuehao=dianming.xuehao and dianming.xuehao='"+xuehaoKaoqin+"'"
            temp = Common.getDataBySQL(sql)            
            
            for iii, student in enumerate(temp):
                treeXueshengMingdan.insert('', iii, values=(student[0], student[1], student[2]))
        buttonXuehao = tkinter.Button(self.top, text='依學號查詢', command=chakanXuehao)
        buttonXuehao.place(x=160, y=60, height=20, width=80)

        # 建立表格，設定表頭，show="headings"用來隱藏樹狀元件的預設首行
        self.frame = tkinter.Frame(self.top)
        self.frame.place(x=20, y=90, width=290, height=280)
        # 垂直捲軸
        scrollBar = tkinter.Scrollbar(self.frame)
        scrollBar.pack(side=tkinter.RIGHT, fill=tkinter.Y)
        # 以樹狀元件實作表格
        treeXueshengMingdan = tkinter.ttk.Treeview(self.frame,
                                                   columns=('col1', 'col2', 'col3'),
                                                   show="headings",
                                                   yscrollcommand = scrollBar.set)
        treeXueshengMingdan.column('col1', width=70, anchor='center')
        treeXueshengMingdan.column('col2', width=60, anchor='center')
        treeXueshengMingdan.column('col3', width=140, anchor='center')
        treeXueshengMingdan.heading('col1', text='學號')
        treeXueshengMingdan.heading('col2', text='姓名')
        treeXueshengMingdan.heading('col3', text='出勤時間')
        treeXueshengMingdan.pack(side=tkinter.LEFT, fill=tkinter.Y)
        # 結合樹狀元件與垂直捲軸
        scrollBar.config(command=treeXueshengMingdan.yview)
def buttonChakanKaoqinXinxiClick():
    # 如果還沒有註冊，拒絕執行
    if int_zhuce.get() == 0:
        tkinter.messagebox.showerror('很抱歉', '請聯繫作者進行軟體註冊！')
        return
    if int_windowChakanKaoqinXinxi.get()==0:
        int_windowChakanKaoqinXinxi.set(1)
        w1 = windowChakanKaoqinXinxi(root, '查詢出勤情況')
        buttonChakanKaoqinXinxi.wait_window(w1.top)
        int_windowChakanKaoqinXinxi.set(0)
buttonChakanKaoqinXinxi = tkinter.Button(root, text='查詢出勤情況', command=buttonChakanKaoqinXinxiClick)
buttonChakanKaoqinXinxi.place(x=240, y=60, height=30, width=100)
## =====================查詢學生出勤情況功能程式碼結束==============================


## =====================隨機提問功能程式碼開始==============================
int_windowTiwen = tkinter.IntVar(root, value=0)
studentInformation = tkinter.StringVar('')
class windowTiwen:
    def __init__(self, root, myTitle, already=[]):
        '''already表示已提問的學生名單
           這裡充分利用清單作為函數預設值參數時的特性'''
        self.top = tkinter.Toplevel(root, width=300, height=150)
        self.top.title(myTitle)
        self.top.attributes('-topmost', 1)

        #學生專業
        xueshengZhuanye = Common.getZhuanye()
        comboboxZhuanye = tkinter.ttk.Combobox(self.top, values=xueshengZhuanye)
        comboboxZhuanye.place(x=20, y=20, height=20, width=100)
        
        #被提問到的學生學號
        xueshengXuehao = tkinter.StringVar(self.top, value='')
        
        def buttonTiwenClick():
            zhuanye = comboboxZhuanye.get()
            if not zhuanye:
                tkinter.messagebox.showerror(title='很抱歉', message='請選擇一個專業！')
                return
            # 從該課程學生名單中隨機選取一個
            conn = sqlite3.connect('database.db')
            cur = conn.cursor()
            cur.execute("select xuehao,xingming from students where zhuanye='"+zhuanye+"'")
            allStudents = cur.fetchall()
            conn.close
            # 位何同學不會在同一節課被提問兩次
            while True:
                random.shuffle(allStudents)
                random.shuffle(allStudents)
                temp = random.choice(allStudents)
                if temp not in already:
                    already.append(temp)
                    #print(already)
                    break
            xuehao = temp[0]
            xueshengXuehao.set(xuehao)
            studentInformation.set('本次中獎同學為：'+str((temp[0],temp[1])))
            #tkinter.messagebox.showinfo(title='恭喜', message='本次中獎同學為'+str((temp[0],temp[1])))
        
        buttonTiwen = tkinter.Button(self.top, text='看看誰最幸運', command=buttonTiwenClick)
        buttonTiwen.place(x=130, y=20, height=20, width=80)
        comboboxDefen = tkinter.ttk.Combobox(self.top, values=['-2', '-1', '0', '1', '2', '3', '4', '5'])
        comboboxDefen.place(x=20, y=50, height=20, width=100)
        def buttonDefenClick():
            if xueshengXuehao.get()=='':
                tkinter.messagebox.showerror(title='很抱歉', message='請挑選一位同學！')
                return
            defen = comboboxDefen.get()
            if not defen:
                tkinter.messagebox.showerror(title='很抱歉', message='請選擇得分！')
            else:
                # 記錄該學生得分
                sql = "insert into tiwen(xuehao,shijian,defen) values('"+xueshengXuehao.get()+"','"+Common.getCurrentDateTime()+"',"+defen+")"
                Common.doSQL(sql)
                if int(defen) > 0:
                    tkinter.messagebox.showinfo(title='太棒了', message='加分成功！')
                elif int(defen) < 0:
                    tkinter.messagebox.showerror('太慘了', message='被減分了！！！')
                
        buttonDefen = tkinter.Button(self.top, text='確認得分', command=buttonDefenClick)
        buttonDefen.place(x=130, y=50, height=20, width=80)

        lbInformation = tkinter.Label(self.top, fg='red', textvariable=studentInformation)
        lbInformation.place(x=20, y=100, height=20, width=240)
def buttonTiwenClick():
    # 如果還沒有註冊，拒絕執行
    if int_zhuce.get() == 0:
        tkinter.messagebox.showerror('很抱歉', '請聯繫作者進行軟體註冊！')
        return
    if int_windowTiwen.get()==0:
        int_windowTiwen.set(1)
        w1 = windowTiwen(root, '隨機提問')
        buttonTiwen.wait_window(w1.top)
        int_windowTiwen.set(0)
buttonTiwen = tkinter.Button(root, text='隨機提問', command=buttonTiwenClick)
buttonTiwen.place(x=20, y=100, height=30, width=100)
## =====================隨機提問功能程式碼結束==============================


## =====================查詢提問情況功能程式碼開始==============================
int_windowChakanTiwenQingkuang = tkinter.IntVar(root, value=0)
class windowChakanTiwenQingkuang:
    def __init__(self, root, myTitle):
        self.top = tkinter.Toplevel(root, width=320, height=380)
        self.top.title(myTitle)
        self.top.attributes('-topmost', 1)
        # 下拉式清單選擇專業，或文字框查詢特定學生的提問情況
        xueshengZhuanye = Common.getZhuanye()
        comboboxZhuanye = tkinter.ttk.Combobox(self.top, values=xueshengZhuanye)
        def comboboxZhuanyeChange(event):
            zhuanye = comboboxZhuanye.get()
            if zhuanye:
                xueshengs = Common.getXuehaoXingming(zhuanye)
                comboboxXuehao['values'] = xueshengs                
        comboboxZhuanye.bind('<<ComboboxSelected>>', comboboxZhuanyeChange)
        comboboxZhuanye.place(x=20, y=20, height=20, width=120)

        #查詢指定專業所有同學的提問情況
        def chakanZhuanye():
            zhuanye = comboboxZhuanye.get()
            if not zhuanye:
                tkinter.messagebox.showerror('很抱歉', '請選擇一個專業！')
                return
            else:
                conn = sqlite3.connect('database.db')
                cur = conn.cursor()
                cur.execute("select students.xuehao,students.xingming,tiwen.shijian,tiwen.defen from students,tiwen where students.xuehao=tiwen.xuehao and students.zhuanye='"+zhuanye+"' order by students.xuehao, tiwen.shijian desc")
                temp = cur.fetchall()
                conn.close

                # 刪除原有的所有列
                for row in treeXueshengMingdan.get_children():
                    treeXueshengMingdan.delete(row)

                for iii, student in enumerate(temp):
                    treeXueshengMingdan.insert('', iii, values=(student[0], student[1], student[2], student[3]))

        buttonZhuanye = tkinter.Button(self.top, text='依專業查詢', command=chakanZhuanye)
        buttonZhuanye.place(x=150, y=20, height=20, width=80)
        xueshengXuehao =[]
        comboboxXuehao = tkinter.ttk.Combobox(self.top, values=xueshengXuehao)
        comboboxXuehao.place(x=20, y=60, height=20, width=120)
        
        def chakanXuehao():
            # 下拉式清單是"學號,姓名"的格式，必須取出學號資料
            xuehao = comboboxXuehao.get().split(',')[0]
            if xuehao=='':
                tkinter.messagebox.showerror('很抱歉', '請選擇學號！')
                return
            else:
                conn = sqlite3.connect('database.db')
                cur = conn.cursor()
                cur.execute("select students.xuehao,students.xingming,tiwen.shijian,tiwen.defen from students,tiwen where students.xuehao=tiwen.xuehao and students.xuehao='"+xuehao+"' order by tiwen.shijian desc")
                temp = cur.fetchall()
                conn.close
                
                #刪除原有的所有列
                for row in treeXueshengMingdan.get_children():
                    treeXueshengMingdan.delete(row)
                
                for iii, student in enumerate(temp):
                    treeXueshengMingdan.insert('', iii, values=(student[0], student[1], student[2], student[3]))
        buttonXuehao = tkinter.Button(self.top, text='依學號查詢', command=chakanXuehao)
        buttonXuehao.place(x=150, y=60, height=20, width=80)
        
        # 建立表格，設定表頭，show="headings"用來隱藏樹狀元件的預設首行
        self.frame = tkinter.Frame(self.top)
        self.frame.place(x=20, y=90, width=300, height=280)
        # 垂直捲軸
        scrollBar = tkinter.Scrollbar(self.frame)
        scrollBar.pack(side=tkinter.RIGHT, fill=tkinter.Y)
        # 以樹狀元件實作表格
        treeXueshengMingdan = tkinter.ttk.Treeview(self.frame,
                                                   columns=('col1', 'col2', 'col3', 'col4'),
                                                   show="headings",
                                                   yscrollcommand = scrollBar.set)
        treeXueshengMingdan.column('col1', width=70, anchor='center')
        treeXueshengMingdan.column('col2', width=50, anchor='center')
        treeXueshengMingdan.column('col3', width=120, anchor='center')
        treeXueshengMingdan.column('col4', width=40, anchor='center')
        treeXueshengMingdan.heading('col1', text='學號')
        treeXueshengMingdan.heading('col2', text='姓名')
        treeXueshengMingdan.heading('col3', text='提問時間')
        treeXueshengMingdan.heading('col4', text='得分')
        treeXueshengMingdan.pack(side=tkinter.LEFT, fill=tkinter.Y)
        # 結合樹狀元件與垂直捲軸
        scrollBar.config(command=treeXueshengMingdan.yview)
        
def buttonChakanTiwenQingkuangClick():
    # 如果還沒有註冊，拒絕執行
    if int_zhuce.get() == 0:
        tkinter.messagebox.showerror('很抱歉', '請聯繫作者進行軟體註冊！')
        return
    if int_windowChakanTiwenQingkuang.get()==0:
        int_windowChakanTiwenQingkuang.set(1)
        w1 = windowChakanTiwenQingkuang(root, '查詢提問情況')
        buttonChakanTiwenQingkuang.wait_window(w1.top)
        int_windowChakanTiwenQingkuang.set(0)
buttonChakanTiwenQingkuang = tkinter.Button(root, text='查詢提問情況', command=buttonChakanTiwenQingkuangClick)
buttonChakanTiwenQingkuang.place(x=130, y=100, height=30, width=100)
## =====================查詢提問情況功能程式碼結束==============================

## =====================查詢統計情況功能程式碼開始==============================
int_windowChakanTongjiQingkuang = tkinter.IntVar(root, value=0)
class windowChakanTongjiQingkuang:
    def __init__(self, root, myTitle):
        self.top = tkinter.Toplevel(root, width=600, height=380)
        self.top.title(myTitle)
        self.top.attributes('-topmost', 1)
        # 下拉式清單選擇專業，或文字框查詢特定學生的提問情況
        xueshengZhuanye = Common.getZhuanye()  #學生專業
        comboboxZhuanye = tkinter.ttk.Combobox(self.top, values=xueshengZhuanye)
        comboboxZhuanye.place(x=20, y=20, height=20, width=120)

        #查詢指定專業所有同學的提問情況
        def chakanZhuanye():
            zhuanye = comboboxZhuanye.get()
            if not zhuanye:
                tkinter.messagebox.showerror('很抱歉', '請選擇一個專業！')
                return
            else:
                xuehaoXingmings = Common.getXuehaoXingming(zhuanye)
                xuehaos = [xingming.split(',')[0] for xingming in xuehaoXingmings]
                xingmings = [xingming.split(',')[1] for xingming in xuehaoXingmings]
                
                # 抓取每位同學的出勤次數，缺勤算0，沒提問到也算0
                chuqinCishu = [Common.getChuqinCishu(xuehao) for xuehao in xuehaos]
                tiwenDefen = [Common.getTiwenDefen(xuehao) for xuehao in xuehaos]
                zhudongTiwenCishu = [Common.getZhudongTiwenCishu(xuehao) for xuehao in xuehaos]
                kaoshidefen = [Common.getKaoshiDefen(xuehao) for xuehao in xuehaos]

                #刪除原有的所有列
                for row in treeXueshengMingdan.get_children():
                    treeXueshengMingdan.delete(row)
                iii = 0
                for xuehao, xingming, chuqin, tiwen, zhudongtiwen, kaoshidefen in zip(xuehaos, xingmings, chuqinCishu, tiwenDefen, zhudongTiwenCishu, kaoshidefen):
                    treeXueshengMingdan.insert('', iii, values=(xuehao, xingming, chuqin, tiwen, zhudongtiwen, kaoshidefen))
                    iii = iii+1
        buttonZhuanye = tkinter.Button(self.top, text='依專業查詢', command=chakanZhuanye)
        buttonZhuanye.place(x=150, y=20, height=20, width=80)
        
        # 建立表格，設定表頭，show="headings"用來隱藏樹狀元件的預設首行
        self.frame = tkinter.Frame(self.top)
        self.frame.place(x=20, y=50, width=560, height=320)
        # 垂直捲軸
        scrollBar = tkinter.Scrollbar(self.frame)
        scrollBar.pack(side=tkinter.RIGHT, fill=tkinter.Y)
        # 以樹狀元件實作表格
        treeXueshengMingdan = tkinter.ttk.Treeview(self.frame,
                                                   columns=('col1', 'col2', 'col3', 'col4', 'col5', 'col6'),
                                                   show="headings",
                                                   yscrollcommand = scrollBar.set)
        treeXueshengMingdan.column('col1', width=70, anchor='center')
        treeXueshengMingdan.column('col2', width=50, anchor='center')
        treeXueshengMingdan.column('col3', width=120, anchor='center')
        treeXueshengMingdan.column('col4', width=120, anchor='center')
        treeXueshengMingdan.column('col5', width=80, anchor='center')
        treeXueshengMingdan.column('col6', width=80, anchor='center')
        treeXueshengMingdan.heading('col1', text='學號')
        treeXueshengMingdan.heading('col2', text='姓名')
        treeXueshengMingdan.heading('col3', text='出勤次數')
        treeXueshengMingdan.heading('col4', text='老師提問得分')
        treeXueshengMingdan.heading('col5', text='主動提問次數')
        treeXueshengMingdan.heading('col6', text='考試得分')
        treeXueshengMingdan.pack(side=tkinter.LEFT, fill=tkinter.Y)
        # 結合樹狀元件與垂直捲軸
        scrollBar.config(command=treeXueshengMingdan.yview)
        
def buttonChakanTongjiQingkuangClick():
    # 如果還沒有註冊，拒絕執行
    if int_zhuce.get() == 0:
        tkinter.messagebox.showerror('很抱歉', '請聯繫作者進行軟體註冊！')
        return
    if int_windowChakanTongjiQingkuang.get()==0:
        int_windowChakanTongjiQingkuang.set(1)
        w1 = windowChakanTongjiQingkuang(root, '查詢統計情況')
        buttonChakanTongjiQingkuang.wait_window(w1.top)
        int_windowChakanTongjiQingkuang.set(0)
buttonChakanTongjiQingkuang = tkinter.Button(root, text='查詢統計情況', command=buttonChakanTongjiQingkuangClick)
buttonChakanTongjiQingkuang.place(x=240, y=100, height=30, width=100)
## =====================查詢統計情況功能程式碼結束==============================


## =====================遠端截圖交作業功能程式碼開始==============================
# 0停止收作業，1開始收作業
int_zuoye = tkinter.IntVar(root, value=0)

def thread_ShouZuoye(conn, today):
    # 最後把客戶端傳來的截圖儲存為"學號，姓名.jpg"
    # 開始接收客戶端送來的截圖
    BUFSIZE = 1024
    FILEINFO_SIZE=struct.calcsize('I128sI')
    fhead = conn.recv(FILEINFO_SIZE)
    filenamelength, filename,filesize=struct.unpack('I128sI',fhead)
    filename = filename.decode()
    filename = filename[:filenamelength]
    ttt = filename.split('.')[0]
    
    # 如果學號和姓名不符合，拒絕接收作業圖片
    xuehao, xingming = ttt.split('_')
    xuehao = xuehao.replace(';', '').replace('"', '').replace("'", '').replace('=', '')
    xingming = xingming.replace(';', '').replace('"', '').replace("'", '').replace('=', '')
    sql = "select count(xuehao) from students where xuehao='"+xuehao.strip()+"' and xingming='"+xingming.strip()+"'"
    t = Common.getDataBySQL(sql)[0][0]
    if t != 1:
        conn.sendall('notmatch'.encode())
        conn.close()
        return
    else:
        conn.sendall('ok'.encode())
    filename = filename[:-4]+'_'.join(Common.getCurrentDateTime().split())+filename[-4:]
    filename = filename.replace('-', '_')
    filename = filename.replace(':', '_')
    filename = today+'\\'+filename

    # 首先刪除本次作業期間之前上交的作業，只保留最後一次的作業
    for f in os.listdir(today):
        if f.startswith(ttt):
            os.remove(today+'\\'+f)
    
    fp = open(filename,'wb')
    restsize = filesize
    while True:
        if restsize > BUFSIZE:
            filedata = conn.recv(BUFSIZE)
        else:
            filedata = conn.recv(restsize)
        if not filedata:
            break
        fp.write(filedata)
        restsize = restsize-len(filedata)
        if restsize == 0:
            break
    fp.close()
    conn.close()
    
def thread_ZuoyeMain():
    today = Common.getCurrentDateTime().split()[0]
    if not os.path.exists(today):
        os.mkdir(today)
    # 開始監聽
    global sockShouzuoye
    sockShouzuoye = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    sockShouzuoye.bind(('', 30100))
    sockShouzuoye.listen(200)
    while int_zuoye.get() == 1:
        time.sleep(0.05)
        try:
            conn, addr = sockShouzuoye.accept()
        except:
            return
        t = threading.Thread(target=thread_ShouZuoye, args=(conn,today))
        t.start()
    sock.close()
        
def buttonStartShouzuoyeClick():
    # 如果還沒有註冊，拒絕執行
    if int_zhuce.get() == 0:
        tkinter.messagebox.showerror('很抱歉', '請聯繫作者進行軟體註冊！')
        return
    if int_canDianming.get() == 1:
        tkinter.messagebox.showerror('很抱歉', '現在正在點名！')
        return
    if int_zuoye.get() == 1:
        tkinter.messagebox.showerror('很抱歉', '現在正在收作業，不用重複設定！')
        return
    int_zuoye.set(1)
    tkinter.messagebox.showinfo('恭喜', '設定成功，現在開始收作業！')
    t = threading.Thread(target=thread_ZuoyeMain)
    # 啟動收作業執行緒
    t.start()
buttonStartShouzuoye = tkinter.Button(root, text='開始收截圖作業', command=buttonStartShouzuoyeClick)
buttonStartShouzuoye.place(x=20, y=140, height=30, width=100)

def buttonStopShouzuoyeClick():
    # 如果還沒有註冊，拒絕執行
    if int_zhuce.get() == 0:
        tkinter.messagebox.showerror('很抱歉', '請聯繫作者進行軟體註冊！')
        return
    if int_zuoye.get()==0:
        tkinter.messagebox.showerror('很抱歉', '還沒開始收作業！')
        return
    int_zuoye.set(0)
    sockShouzuoye.close()
    time.sleep(0.1)
    tkinter.messagebox.showinfo('恭喜', '設定成功，現在已停止收作業！')
buttonStopShouzuoye = tkinter.Button(root, text='結束收截圖作業', command=buttonStopShouzuoyeClick)
buttonStopShouzuoye.place(x=130, y=140, height=30, width=100)
## =====================遠端截圖交作業功能程式碼結束==============================

## =====================聯繫方式與註冊功能程式碼開始==============================
# 聯繫作者，註冊功能暫時不用，軟體免費試用
int_windowLianxiFangshi = tkinter.IntVar(root, value=0)
class windowLianxiFangshi:
    def __init__(self, root, myTitle):
        self.top = tkinter.Toplevel(root, width=420, height=200)
        self.top.title(myTitle)
        self.top.attributes('-topmost', 1)
        labelQQ = tkinter.Label(self.top, text='QQ:306467355', justify=tkinter.LEFT, width=100)
        labelQQ.place(x=20, y=10, height=20, width=100)
        labelEmail = tkinter.Label(self.top, text='Email:dongfuguo2005@126.com', justify=tkinter.LEFT, width=200)
        labelEmail.place(x=20, y=40, height=20, width=200)
        labelWeixin = tkinter.Label(self.top, text='微信:Python_dfg', justify=tkinter.LEFT, width=110)
        labelWeixin.place(x=20, y=70, height=20, width=100)
        
        # 抓取本機網卡實際位址
        import uuid
        node = uuid.getnode()
        macHex = uuid.UUID(int=node).hex[-12:]
        self.mac = []
        for i in range(len(macHex))[::2]:
            self.mac.append(macHex[i:i+2])
        self.mac = ''.join(self.mac)
        import string
        tttt = string.ascii_lowercase+string.digits
        mac1 = ''.join([random.choice(tttt) for i in range(10)])+self.mac+''.join([random.choice(tttt) for i in range(10)])
        string_mac = tkinter.StringVar(self.top)
        string_mac.set(mac1)
        entryMac = tkinter.Entry(self.top, textvariable=string_mac)
        entryMac.place(x=10, y=110, height=20, width=400)
        
        string_zhucema = tkinter.StringVar(self.top)
        entryZhucema = tkinter.Entry(self.top, textvariable=string_zhucema)
        entryZhucema.place(x=10, y=140, height=20, width=400)
        
        # 註冊按鈕功能
        def buttonZhuceClick():
            # 計算註冊碼
            import hashlib
            zhucema = hashlib.md5(self.mac.encode()).hexdigest()
            # 抓取使用者輸入的註冊碼
            zhucema1 = entryZhucema.get()
            if len(zhucema1.strip())<60 or zhucema != zhucema1[15:47]:
                tkinter.messagebox.showerror('很抱歉', '註冊碼不正確！')
                return
            else:
                # 產生註冊檔案
                tttt = string.ascii_lowercase+string.digits
                fileContent = ''.join([random.choice(tttt) for i in range(3000)])
                fileContent = fileContent[:35]+'8xm0a'+fileContent[40:]
                fileContent = fileContent[0:150]+zhucema+fileContent[182:]
                with open('kaoqin.txt', 'w') as fp:
                    fp.write(fileContent)                
                tkinter.messagebox.showinfo('恭喜', '註冊成功！')
                isZhuce()
                self.top.destroy()
        buttonZhuce = tkinter.Button(self.top, text='註冊', command=buttonZhuceClick)
        buttonZhuce['state'] = 'disabled'
        buttonZhuce.place(x=80, y=170, height=20, width=100)
        
def buttonLianxiFangshiClick():
    if int_windowLianxiFangshi.get()==0:
        int_windowLianxiFangshi.set(1)
        w1 = windowLianxiFangshi(root, '作者聯繫方式')
        buttonChakanTongjiQingkuang.wait_window(w1.top)
        int_windowLianxiFangshi.set(0)
buttonLianxiFangshi = tkinter.Button(root, text='聯繫作者', command=buttonLianxiFangshiClick)
buttonLianxiFangshi.place(x=240, y=140, height=30, width=100)
## =====================聯繫方式與註冊功能程式碼結束==============================

## =====================線上答問功能程式碼開始==============================
# 程式啟動後立刻啟動一個執行緒，隨時接收學生提問，彈出聯天對話框
# 關閉聊天對話框時把聊天內容寫入資料庫
# 控制學生是否可以提問的變數
int_xueshengTiwen = tkinter.IntVar(root, value=0)
class windowXueshengTiwen:
    def __init__(self, root, conn, data):
        # 建立面板容器，用來擺放其他元件
        self.top = tkinter.Toplevel(root, width=300, height=220)
        tttt = tkinter.StringVar(self.top)
        xuehaoXingming, message = data.split(':')
        tttt.set(message)
        self.top.title('來自'+xuehaoXingming+' 的消息')
        # 將本次學生提問記錄寫到資料庫
        sqlTiwen = "insert into xueshengtiwen(xuehao,wenti,shijian) values('"\
                   +xuehaoXingming.split(',')[0]+"','"+message+"','"+Common.getCurrentDateTime()+"')"
        Common.doSQL(sqlTiwen)
        # 顯示視窗
        self.top.attributes('-topmost', 1)
        # 接收到的訊息
        entryMessage = tkinter.scrolledtext.ScrolledText(self.top, wrap=tkinter.WORD)
        entryMessage.insert(tkinter.INSERT, message)
        entryMessage.place(x=20, y=20, width=250, height=60)
        # 要回覆的訊息
        entryHuifu = tkinter.scrolledtext.ScrolledText(self.top, wrap=tkinter.WORD)
        entryHuifu.place(x=20, y=90, width=250, height=60)
        
        def buttonHuifuClick():
            # 回覆，解答
            huifu = entryHuifu.get(0.0, tkinter.END)
            conn.sendall(huifu.encode())
            conn.close()
            # 將本次回覆記錄寫入資料庫
            sqlTiwen = "insert into xueshengtiwen(xuehao,wenti,shijian) values('"\
                       +xuehaoXingming.split(',')[0]+"','老師回覆 "+xuehaoXingming+":"+huifu\
                       +"','"+Common.getCurrentDateTime()+"')"
            Common.doSQL(sqlTiwen)
            # 關閉視窗
            self.top.destroy()
        buttonHuifu = tkinter.Button(self.top, text='回覆解答', command=buttonHuifuClick)
        buttonHuifu.place(x=40, y=160, width=80, height=20)

        def buttonTongyiJiangjieClick():
            # 不回覆該問題，稍後統一講解給全班同學
            conn.sendall('wait'.encode())
            conn.close()
            # 將本次回覆記錄寫入資料庫
            sqlTiwen = "insert into xueshengtiwen(xuehao,wenti,shijian) values('"\
                       +xuehaoXingming.split(',')[0]+"','老師回覆 "+xuehaoXingming+"：該問題稍後統一講解！','"\
                       +Common.getCurrentDateTime()+"')"
            Common.doSQL(sqlTiwen)
            # 關閉視窗
            self.top.destroy()
        buttonTongyiJiangjie = tkinter.Button(self.top, text='稍後統一講解', command=buttonTongyiJiangjieClick)
        buttonTongyiJiangjie.place(x=130, y=160, width=80, height=20)
        
def thread_xueshengTiwenMain():
    global sockTiwen
    sockTiwen = socket.socket(socket.AF_INET , socket.SOCK_STREAM)
    # 監聽20000連接埠
    sockTiwen.bind(('', 20000))
    # 最多允許同時200個學生提問
    sockTiwen.listen(200)
    while int_xueshengTiwen.get() == 1:
        try:
            #接收一個連線
            conn, addr = sockTiwen.accept()
        except:
            return
        #接收訊息，內容為"學號,姓名:問題內容"
        data = conn.recv(1024)
        data = data.decode()
        xuehaoXingming, message = data.split(':')
        xuehao, xingming = xuehaoXingming.split(',')
        xuehao = xuehao.replace(';', '').replace('"', '').replace("'", '').replace('=', '')
        xingming = xingming.replace(';', '').replace('"', '').replace("'", '').replace('=', '')
        sql = "select count(xuehao) from students where xuehao='"+xuehao.strip()+"' and xingming='"+xingming.strip()+"'"
        t = Common.getDataBySQL(sql)[0][0]
        if t != 1:
            conn.sendall('notmatch'.encode())
            conn.close()
        else:
            w = windowXueshengTiwen(root, conn, data)        
    sockTiwen.close()
    
# 啟動接收學生提問的執行緒
def buttonJieshouTiwenClick():
    # 如果還沒有註冊，拒絕執行
    if int_zhuce.get() == 0:
        tkinter.messagebox.showerror('很抱歉', '請聯繫作者進行軟體註冊！')
        return
    if int_xueshengTiwen.get() == 1:
        tkinter.messagebox.showerror('很抱歉', '現在正在接收提問，不用重複設定！')
        return
    int_xueshengTiwen.set(1)
    tkinter.messagebox.showinfo('恭喜', '設定成功，現在開始接收學生提問！')
    t_tiwen = threading.Thread(target=thread_xueshengTiwenMain)
    t_tiwen.start()
buttonJieshouTiwen = tkinter.Button(root, text='開始接收提問', command=buttonJieshouTiwenClick)
buttonJieshouTiwen.place(x=20, y=180, height=30, width=100)

# 停止接收學生提問的執行緒
def buttonTingzhiTiwenClick():
    # 如果還沒有註冊，拒絕執行
    if int_zhuce.get() == 0:
        tkinter.messagebox.showerror('很抱歉', '請聯繫作者進行軟體註冊！')
        return
    if int_xueshengTiwen.get() == 0:
        tkinter.messagebox.showerror('很抱歉', '還沒開始接收學生提問！')
        return
    int_xueshengTiwen.set(0)
    sockTiwen.close()
    time.sleep(0.1)
    tkinter.messagebox.showinfo('恭喜', '設定成功，現在停止接收學生提問！')
buttonTingzhiTiwen = tkinter.Button(root, text='停止接收提問', command=buttonTingzhiTiwenClick)
buttonTingzhiTiwen.place(x=130, y=180, height=30, width=100)
## =====================線上答問功能程式碼結束==============================

## =====================查詢學生主動提問情況功能程式碼開始==============================
int_windowXueshengTiwenQingkuang = tkinter.IntVar(root, value=0)
class windowXueshengTiwenQingkuang:
    def __init__(self, root, myTitle):
        self.top = tkinter.Toplevel(root, width=620, height=380)
        self.top.title(myTitle)
        self.top.attributes('-topmost', 1)
        # 下拉式清單選擇專業，或文字框查詢特定學生的提問情況
        xueshengZhuanye = Common.getZhuanye()
        comboboxZhuanye = tkinter.ttk.Combobox(self.top, values=xueshengZhuanye)
        
        def comboboxZhuanyeChange(event):
            zhuanye = comboboxZhuanye.get()
            if zhuanye:
                xueshengs = Common.getXuehaoXingming(zhuanye)
                comboboxXuehao['values'] = xueshengs                
        comboboxZhuanye.bind('<<ComboboxSelected>>', comboboxZhuanyeChange)
        comboboxZhuanye.place(x=20, y=20, height=20, width=120)

        # 查詢指定專業所有同學的提問情況
        def chakanZhuanye():
            zhuanye = comboboxZhuanye.get()
            if not zhuanye:
                tkinter.messagebox.showerror('很抱歉', '請選擇一個專業！')
                return
            else:
                conn = sqlite3.connect('database.db')
                cur = conn.cursor()
                cur.execute("select students.xuehao,students.xingming,xueshengtiwen.wenti,xueshengtiwen.shijian from students,xueshengtiwen where students.xuehao=xueshengtiwen.xuehao and students.zhuanye='"+zhuanye+"' order by students.xuehao, xueshengtiwen.shijian desc")
                temp = cur.fetchall()
                conn.close
                # 刪除原有的所有列
                for row in treeXueshengMingdan.get_children():
                    treeXueshengMingdan.delete(row)

                for iii, student in enumerate(temp):
                    treeXueshengMingdan.insert('', iii, values=(student[0], student[1], student[2], student[3]))
        buttonZhuanye = tkinter.Button(self.top, text='依專業查詢', command=chakanZhuanye)
        buttonZhuanye.place(x=150, y=20, height=20, width=80)
        xueshengXuehao =[]
        comboboxXuehao = tkinter.ttk.Combobox(self.top, values=xueshengXuehao)
        comboboxXuehao.place(x=20, y=60, height=20, width=120)
        
        def chakanXuehao():
            # 下拉式清單是"學號,姓名"的格式，必須取出學號資料
            xuehao = comboboxXuehao.get().split(',')[0]
            if xuehao=='':
                tkinter.messagebox.showerror('很抱歉', '請選擇學號！')
                return
            else:
                conn = sqlite3.connect('database.db')
                cur = conn.cursor()
                cur.execute("select students.xuehao,students.xingming,xueshengtiwen.wenti,xueshengtiwen.shijian from students,xueshengtiwen where students.xuehao=xueshengtiwen.xuehao and students.xuehao='"+xuehao+"' order by xueshengtiwen.shijian desc")
                temp = cur.fetchall()
                conn.close
                # 刪除原有的所有列
                for row in treeXueshengMingdan.get_children():
                    treeXueshengMingdan.delete(row)

                for iii, student in enumerate(temp):
                    treeXueshengMingdan.insert('', iii, values=(student[0], student[1], student[2], student[3]))
        buttonXuehao = tkinter.Button(self.top, text='依學號查詢', command=chakanXuehao)
        buttonXuehao.place(x=150, y=60, height=20, width=80)
        
        # 建立表格，設定表頭，show="headings"用來隱藏樹狀元件的預設首行
        self.frame = tkinter.Frame(self.top)
        self.frame.place(x=10, y=90, width=600, height=280)
        # 垂直捲軸
        scrollBar = tkinter.Scrollbar(self.frame)
        scrollBar.pack(side=tkinter.RIGHT, fill=tkinter.Y)
        # 以樹狀元件實作表格
        treeXueshengMingdan = tkinter.ttk.Treeview(self.frame,
                                                   columns=('col1', 'col2', 'col3', 'col4'),
                                                   show="headings",
                                                   yscrollcommand = scrollBar.set)
        treeXueshengMingdan.column('col1', width=70, anchor='center')
        treeXueshengMingdan.column('col2', width=50, anchor='center')
        treeXueshengMingdan.column('col3', width=320, anchor='center')
        treeXueshengMingdan.column('col4', width=140, anchor='center')
        treeXueshengMingdan.heading('col1', text='學號')
        treeXueshengMingdan.heading('col2', text='姓名')
        treeXueshengMingdan.heading('col3', text='問題')
        treeXueshengMingdan.heading('col4', text='時間')
        treeXueshengMingdan.pack(side=tkinter.LEFT, fill=tkinter.Y)
        # 結合樹狀元件與垂直捲軸
        scrollBar.config(command=treeXueshengMingdan.yview)
        
def buttonChakanTiwenQingkuangClick():
    # 如果還沒有註冊，拒絕執行
    if int_zhuce.get() == 0:
        tkinter.messagebox.showerror('很抱歉', '請聯繫作者進行軟體註冊！')
        return
    if int_windowXueshengTiwenQingkuang.get()==0:
        int_windowXueshengTiwenQingkuang.set(1)
        w1 = windowXueshengTiwenQingkuang(root, '學生主動提問情況')
        buttonXueshengTiwenQingkuang.wait_window(w1.top)
        int_windowXueshengTiwenQingkuang.set(0)
buttonXueshengTiwenQingkuang = tkinter.Button(root, text='學生主動提問情況', command=buttonChakanTiwenQingkuangClick)
buttonXueshengTiwenQingkuang.place(x=240, y=180, height=30, width=100)
## =====================查詢學生主動提問情況功能程式碼結束==============================


## =====================題庫匯入功能程式碼開始==============================
def buttonDaoruTikuClick():
    # 如果還有沒註冊，拒絕執行
    if int_zhuce.get() == 0:
        tkinter.messagebox.showerror('很抱歉', '請聯繫作者進行軟體註冊！')
        return
    filename = tkinter.filedialog.askopenfilename(title='請選擇Excel 2007-或Word 2007+版本的題庫檔',
                                                          filetypes=[('Excel Files','*.xls'), ('Word 2007+ Files','*.docx')])
    if filename:
        # 讀取資料並匯入資料庫
        # Excel檔案
        if filename.endswith('.xls'):
            workbook = xlrd.open_workbook(filename=filename)
            sheet1 = workbook.sheet_by_index(0)
            # Excel檔案必須包含4行，分別是課程名稱、章節、題目、答案
            if sheet1.ncols != 4:
                tkinter.messagebox.showerror(title='很抱歉', message='Excel檔案格式不對！')
                return

            # 巡訪Excel檔案每一列
            for rowIndex in range(1, sheet1.nrows):
                row = sheet1.row(rowIndex)
                sql = "insert into tiku(kechengmingcheng,zhangjie,timu,daan) values('"\
                      + "','".join(map(lambda item:str(item.value).strip(), row)) + "')"
                Common.doSQL(sql)
            tkinter.messagebox.showinfo(title='恭喜', message='匯入成功')
            
        # docx檔案
        elif filename.endswith('.docx'):
            # docx檔案題庫包含很多段，每段一個題目，格式為：   問題。（答案）
            # 資料庫datase.db中tiku資料表包含kechengmingcheng,zhangjie,timu,daan四個欄位
            from docx import Document
            doc = Document(filename)

            # 連接資料庫
            conn = sqlite3.connect('database.db')
            cur = conn.cursor()

            # 先清空原來的題目，可有可無
            cur.execute('delete from tiku')
            conn.commit()

            for p in doc.paragraphs:
                text = p.text
                if '（' in text and '）' in text:
                    index = text.index('（')
                    # 分離問題和答案
                    question = text[:index]
                    if '___' in question:
                        question = '填空題：' + question
                    else:
                        question = '是非題：' + question
                    answer = text[index+1:-1]
                    # 將資料寫入資料庫
                    sql = 'insert into tiku(kechengmingcheng,zhangjie,timu,daan) values("Python程式設計","未分類","'\
                          +question+'","'+answer+'")'
                    cur.execute(sql)
            conn.commit()
            # 關閉資料庫連線
            conn.close()
            tkinter.messagebox.showinfo(title='恭喜', message='匯入成功')
buttonDaoruTiku = tkinter.Button(root, text='匯入題庫', command=buttonDaoruTikuClick)
buttonDaoruTiku.place(x=20, y=220, height=30, width=100)
## =====================題庫匯入功能程式碼結束==============================

## =====================線上自測功能程式碼開始==============================
# 控制學生是否可以自測的變數
int_xueshengZixue = tkinter.IntVar(root, value=0)
def thread_xueshengZixue(conn):
    data = conn.recv(1024)#接收學號和姓名，記錄題庫存取情況
    data = data.decode()
    xuehao, xingming = data.split(',')
    xuehao = xuehao.replace(';', '').replace('"', '').replace("'", '').replace('=', '')
    xingming = xingming.replace(';', '').replace('"', '').replace("'", '').replace('=', '')
    sqlTikuFangwen = "insert into tikufangwenqingkuang(xuehao,xingming,shijian) values('"\
                     +xuehao+"','"+xingming+"','"+Common.getCurrentDateTime()+"')"
    Common.doSQL(sqlTikuFangwen)
    
    # 抓取並向客戶端傳送題庫中的課程名稱，以逗點隔開
    sqlKechengmingcheng = "select distinct(kechengmingcheng) from tiku"
    kechengQingdan = []
    for kecheng in Common.getDataBySQL(sqlKechengmingcheng):
        kechengQingdan.append(str(kecheng[0]))
    kechengQingdan = ','.join(kechengQingdan)
    conn.sendall(kechengQingdan.encode())
    
    # 開始接收課程名稱、目前是號id、上一題/下一題，返回課程名稱、章節、題號id、題目、答案，沒有上一題或下一題則返回no
    # 如果客戶端傳來xxxx，表示結束自測
    while int_xueshengZixue.get() == 1:
        data = conn.recv(1024)
        data = data.decode()
        if data == 'xxxx':
            conn.recv(1024)
            break
        kechengmingcheng, currentID, pre_next = data.split('xx')
        if pre_next == 'pre':#上一題
            sqlHasMore = "select kechengmingcheng,zhangjie,timu,daan,id from tiku where kechengmingcheng='"\
                         +kechengmingcheng+"' and id<"+currentID+" order by id desc"
            ttt = Common.getDataBySQL(sqlHasMore)
            if ttt:
                tttt = ttt[0]
                message = tttt[0]+'xx'+tttt[1]+'xx'+tttt[2]+'xx'+tttt[3]+'xx'+str(tttt[4])
                conn.sendall(message.encode())
            else:
                conn.sendall('no'.encode())
        elif pre_next == 'next':#下一題
            sqlHasMore = "select kechengmingcheng,zhangjie,timu,daan,id from tiku where kechengmingcheng='"\
                         +kechengmingcheng+"' and id>"+currentID+" order by id"
            ttt = Common.getDataBySQL(sqlHasMore)
            if ttt:
                tttt = ttt[0]
                message = tttt[0]+'xx'+tttt[1]+'xx'+tttt[2]+'xx'+tttt[3]+'xx'+str(tttt[4])
                conn.sendall(message.encode())
            else:
                conn.sendall('no'.encode())
    conn.close()
    
def thread_xueshengZixueMain():
    global sockZice
    sockZice = socket.socket(socket.AF_INET , socket.SOCK_STREAM)
    # 監聽10000連接埠
    sockZice.bind(('', 10000))
    # 最多允許同時200位學生提問
    sockZice.listen(200)
    while int_xueshengZixue.get() == 1:
        try:
            # 接收一個連線
            conn, addr = sockZice.accept()
        except:
            return
        t_zixue = threading.Thread(target=thread_xueshengZixue, args=(conn,))
        t_zixue.start()       
    sockZice.close()
    
# 啟動學生自測的執行緒
def buttonStartXueshengZixueClick():
    # 如果還沒有註冊，拒絕執行
    if int_zhuce.get() == 0:
        tkinter.messagebox.showerror('很抱歉', '請聯繫作者進行軟體註冊！')
        return
    if int_xueshengZixue.get() == 1:
        tkinter.messagebox.showerror('很抱歉', '現在正在進行學生自測，不用重複設定！')
        return
    int_xueshengZixue.set(1)
    tkinter.messagebox.showinfo('恭喜', '設定成功，現在開始學生自測！')
    t_tiwen = threading.Thread(target=thread_xueshengZixueMain)
    t_tiwen.start()
buttonStartXueshengZixue = tkinter.Button(root, text='開始學生自測', command=buttonStartXueshengZixueClick)
buttonStartXueshengZixue.place(x=130, y=220, height=30, width=100)

# 停止學生自測的執行緒
def buttonStopXueshengZixueClick():
    # 如果還沒有註冊，拒絕執行
    if int_zhuce.get() == 0:
        tkinter.messagebox.showerror('很抱歉', '請聯繫作者進行軟體註冊！')
        return
    if int_xueshengZixue.get() == 0:
        tkinter.messagebox.showerror('很抱歉', '現在不是學生自測時間！')
        return
    int_xueshengZixue.set(0)
    sockZice.close()
    tkinter.messagebox.showinfo('恭喜', '設定成功，現在停止學生自測！')
buttonStopXueshengZixue = tkinter.Button(root, text='停止學生自測', command=buttonStopXueshengZixueClick)
buttonStopXueshengZixue.place(x=240, y=220, height=30, width=100)
## =====================線上自測功能程式碼結束==============================

## =====================線上上傳檔案交作業功能程式碼開始==============================
int_wenjianZuoye = tkinter.IntVar(root, value=0) #負責控制是否允許接收學生上傳檔案作業的變數
# 負責具體收檔案作業的執行緒函數
def thread_ShouWenjianZuoye(conn, today):
    BUFSIZE = 1024
    FILEINFO_SIZE=struct.calcsize('I128sI')
    fhead = conn.recv(FILEINFO_SIZE)
    filenamelength, filename,filesize=struct.unpack('I128sI',fhead)
    filename = filename.decode()
    filename = filename[:filenamelength]
    
    ttt = filename.split('_1_')[0]
    # 如果學號和姓名不符合，拒絕接收作業
    xuehao, xingming = ttt.split('_')
    xuehao = xuehao.replace(';', '').replace('"', '').replace("'", '').replace('=', '')
    xingming = xingming.replace(';', '').replace('"', '').replace("'", '').replace('=', '')
    sql = "select count(xuehao) from students where xuehao='"+xuehao.strip()+"' and xingming='"+xingming.strip()+"'"
    t = Common.getDataBySQL(sql)[0][0]
    if t != 1:
        conn.sendall('notmatch'.encode())
        conn.close()
        return
    else:
        conn.sendall('ok'.encode())
    index = filename.rindex('.')
    filename = xuehao+'_'+xingming+'_'.join(Common.getCurrentDateTime().split())+filename[index:]
    filename = filename.replace('-', '_')
    filename = filename.replace(':', '_')
    filename = today+'\\'+filename

    # 首先刪除本次作業期間之前上繳的作業，只保留最後一次的作業
    for f in os.listdir(today):
        if f.startswith(ttt):
            os.remove(today+'\\'+f)
    
    fp = open(filename,'wb')
    restsize = filesize
    while True:
        if restsize > BUFSIZE:
            filedata = conn.recv(BUFSIZE)
        else:
            filedata = conn.recv(restsize)
        if not filedata:
            break
        fp.write(filedata)
        restsize = restsize-len(filedata)
        if restsize == 0:
            break
    fp.close()
    conn.close()

# 負責監聽連接埠的收檔案作業主執行緒函數
def thread_wenjianZuoyeMain():
    today = Common.getCurrentDateTime().split()[0]
    if not os.path.exists(today):
        os.mkdir(today)
    # 開始監聽
    global sockShouWenjianZuoye
    sockShouWenjianZuoye = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    sockShouWenjianZuoye.bind(('', 10500))
    sockShouWenjianZuoye.listen(200)
    while int_wenjianZuoye.get() == 1:
        time.sleep(0.05)
        try:
            conn, addr = sockShouWenjianZuoye.accept()
        except:
            return
        t = threading.Thread(target=thread_ShouWenjianZuoye, args=(conn,today))
        t.start()
    sockShouWenjianZuoye.close()
    
# 開始接收學生上傳Python程式檔作業
def buttonStartJieshouWenjianZuoyeClick():
    # 如果還沒有註冊，拒絕執行
    if int_zhuce.get() == 0:
        tkinter.messagebox.showerror('很抱歉', '請聯繫作者進行軟體註冊！')
        return
    if int_wenjianZuoye.get() == 1:
        tkinter.messagebox.showerror('很抱歉', '現在正在收作業，不用重複設定！')
        return
    int_wenjianZuoye.set(1)
    tkinter.messagebox.showinfo('恭喜', '設定成功，現在開始收作業！')
    t = threading.Thread(target=thread_wenjianZuoyeMain)
    #啟動收作業執行緒
    t.start()
buttonStartJieshouWenjianZuoye = tkinter.Button(root,
                                                text='開始收檔案作業',
                                                command=buttonStartJieshouWenjianZuoyeClick)
buttonStartJieshouWenjianZuoye.place(x=20, y=260, height=30, width=100)

# 結束接收學生上傳Python程式檔作業
def buttonStopJieshouWenjianZuoyeClick():
    # 如果還沒有註冊，拒絕執行
    if int_zhuce.get() == 0:
        tkinter.messagebox.showerror('很抱歉', '請聯繫作者進行軟體註冊！')
        return
    if int_wenjianZuoye.get()==0:
        tkinter.messagebox.showerror('很抱歉', '還沒開始收作業！')
        return
    int_wenjianZuoye.set(0)
    sockShouWenjianZuoye.close()
    time.sleep(0.1)
    tkinter.messagebox.showinfo('恭喜', '設定成功，現在已停止收作業！')
buttonStopJieshouWenjianZuoye = tkinter.Button(root,
                                               text='結束收檔案作業',
                                               command=buttonStopJieshouWenjianZuoyeClick)
buttonStopJieshouWenjianZuoye.place(x=130, y=260, height=30, width=100)
## =====================線上上傳檔案交作業功能程式碼結束==============================


## =====================資料匯出功能程式碼開始==============================
# 匯出資料到xlsx檔案
def buttonDaochuClick():
    try:
        import openpyxl
        from openpyxl import Workbook
    except:
        tkinter.messagebox.showerror('抱歉', '您需要安裝openpyxl擴展庫！')
        
    # 設定要匯出資料的時段
    while True:
        startTime = tkinter.simpledialog.askstring('請輸入開始時間', '開始時間：', initialvalue='2016-3-1')
        try:
            import datetime
            # 確保輸入的是正確的日期格式
            y, m, d = map(int, startTime.split('-'))
            startTime = str(datetime.date(y, m, d))
        except:
            pass
        else:
            break

    while True:
        endTime = tkinter.simpledialog.askstring('請輸入結束時間', '結束時間：', initialvalue='2016-3-31')
        try:
            import datetime
            # 確保輸入的是正確的日期格式
            y, m, d = map(int, endTime.split('-'))
            endTime = str(datetime.date(y, m, d))
        except:
            pass
        else:
            break

    # 統一的時間區間條件
    shijianDuan = ' shijian>="' + startTime + '" and shijian<="' + endTime + '"'
    
    wb = Workbook()
    # 刪除預設的worksheet
    wb.remove_sheet(wb.worksheets[0])
    # 點名記錄
    ws = wb.create_sheet(title='線上點名情況')
    ws.append(['學號', '姓名', '點名時間'])
    sql = 'select students.xuehao, students.xingming, shijian from students, dianming where students.xuehao=dianming.xuehao and ' + shijianDuan + ' order by students.xuehao, shijian' 
    data = Common.getDataBySQL(sql)
    for d in data:
        ws.append([d[0], d[1], d[2]])

    # 提問情況
    ws = wb.create_sheet(title='隨機提問與加分情況')
    ws.append(['學號', '姓名', '提問時間', '得分'])
    sql = 'select students.xuehao, students.xingming, shijian, defen from students, tiwen where students.xuehao=tiwen.xuehao and ' + shijianDuan +' order by students.xuehao'
    data = Common.getDataBySQL(sql)
    for d in data:
        ws.append([d[0], d[1], d[2], d[3]])

    # 題庫存取情況
    ws = wb.create_sheet(title='題庫存取情況')
    ws.append(['學號', '姓名', '存取時間'])
    sql = 'select xuehao, xingming, shijian from tikufangwenqingkuang where' + shijianDuan +' order by xuehao '
    data = Common.getDataBySQL(sql)
    for d in data:
        ws.append([d[0], d[1], d[2]])

    # 考試情況
    ws = wb.create_sheet(title='考試情況')
    ws.append(['學號', '姓名', '答題時間', '題目', '標準答案', '學生答案', '是否正確'])
    sql = 'select xuehao, xingming, shijian, timu, daan, xueshengdaan, shifouzhengque from kaoshi,tiku where tiku.id=kaoshi.timubianhao and' + shijianDuan +' order by xuehao, shijian'
    data = Common.getDataBySQL(sql)
    for d in data:
        ws.append([d[0], d[1], d[2], d[3], d[4], d[5], d[6]])

    wb.save('資料匯出.xlsx')
    tkinter.messagebox.showinfo('恭喜', '匯出資料成功，請查看系統資料夾的"資料匯出.xlsx"檔案！')
buttonDaochu = tkinter.Button(root, text='資料匯出', command=buttonDaochuClick)
buttonDaochu.place(x=240, y=260, height=30, width=100)
## =====================資料匯出功能程式碼結束==============================


## =====================線上考試功能程式碼開始==============================
# 控制學生是否可以自測的變數
int_xueshengKaoshi = tkinter.IntVar(root, value=0)
def thread_xueshengKaoshi(conn):
    data = conn.recv(1024)#接收學號和姓名
    data = data.decode()
    xuehao, xingming = data.split(',')
    xuehao = xuehao.replace(';', '').replace('"', '').replace("'", '').replace('=', '')
    xingming = xingming.replace(';', '').replace('"', '').replace("'", '').replace('=', '')
    # 檢查學號、姓名是否符合和正確
    sql = "select count(xuehao) from students where xuehao='"+xuehao.strip()+"' and xingming='"+xingming.strip()+"'"
    t = Common.getDataBySQL(sql)[0][0]
    if t != 1:
        conn.sendall('notmatch'.encode())
        conn.close()
        return
    else:
        conn.sendall('ok'.encode())

    
    sqlTikuFangwen = "insert into tikufangwenqingkuang(xuehao,xingming,shijian) values('"\
                     +xuehao+"','"+xingming+"','"+Common.getCurrentDateTime()+"')"
    Common.doSQL(sqlTikuFangwen)
    # 抓取並向客戶端傳送題庫中的課程名稱，以逗點隔開
    sqlKechengmingcheng = "select distinct(kechengmingcheng) from tiku"
    kechengQingdan = []
    for kecheng in Common.getDataBySQL(sqlKechengmingcheng):
        kechengQingdan.append(str(kecheng[0]))
    kechengQingdan = ','.join(kechengQingdan)
    conn.sendall(kechengQingdan.encode())
    # 開始接收課程名稱、目前題號id、上一題/下一題，返回課程名稱、章節、題號id、題目、答案，沒有上一題或下一題則返回no
    # 如果客戶端傳來xxxx，表示結束考試
    while int_xueshengKaoshi.get() == 1:
        data = conn.recv(1024)
        data = data.decode()
        if data == 'xxxx':
            conn.recv(1024)
            break
        kechengmingcheng, currentID, daan, pre_next = data.split('xx')
        if pre_next == 'next':#下一題
            # 記錄本題學生答案，不記錄0號題
            if currentID != '0':
                if daan == biaozhundaan:
                    # 學生答案正確
                    shifouzhengque = 'Y'
                # 填空題，help和help()兩種形式的答案都算對
                elif biaozhundaan.endswith('()') and daan==biaozhundaan[:-2]:
                    shifouzhengque = 'Y'
                # 如果標準答案中有空格，重新修正學生答案，考慮學生輸入多個連續空格的情況，例如pip list等同於pip   list
                elif ' '.join(biaozhundaan.split()) == ' '.join(daan.split()):
                    shifouzhengque = 'Y'
                elif ''.join(biaozhundaan.split()) == ''.join(daan.split()):
                    #考慮學生未輸入空可的情況，例如[1,2,3]等同於[1, 2, 3]
                    shifouzhengque = 'Y'
                else:
                    shifouzhengque = 'N'
                sql = 'insert into kaoshi(xuehao,xingming,timubianhao,xueshengdaan, biaozhundaan, shifouzhengque, shijian) values("'+xuehao+'","'+xingming+'",'+currentID+',"'+daan+'","'+biaozhundaan+'","'+shifouzhengque+'","'+Common.getCurrentDateTime()+'")'
                Common.doSQL(sql)

            # 判斷學生是否已答100題，若是則不允許繼續答題
            sql = "select count(xuehao) from kaoshi where xuehao='"+xuehao+"'"
            total = Common.getDataBySQL(sql)[0]
            if total[0] >= 100:
                conn.sendall(('no,'+str(Common.getKaoshiDefen(xuehao))).encode())
                break

            # 傳送下一題
            sqlHasMore = "select kechengmingcheng,zhangjie,timu,id, daan from tiku where kechengmingcheng='"+kechengmingcheng+"' and id not in (select timubianhao from kaoshi where xuehao='"+xuehao+"') order by random() limit 1"
            ttt = Common.getDataBySQL(sqlHasMore)
            if ttt:
                tttt = ttt[0]
                message = tttt[0]+'xx'+tttt[1]+'xx'+tttt[2]+'xx'+str(tttt[3])+'xx'+str(total[0])   #total[0]表示已做多少道題
                # 記錄本題答案，評分用
                biaozhundaan = str(tttt[4])
                conn.sendall(message.encode())
            else:
                conn.sendall(('no,'+str(Common.getKaoshiDefen(xuehao))).encode())
    conn.close()
    
def thread_xueshengKaoshiMain():
    global sockKaoshi
    sockKaoshi = socket.socket(socket.AF_INET , socket.SOCK_STREAM)
    # 監聽18001連接埠
    sockKaoshi.bind(('', 18003))
    # 最多允許同時200位學生考試
    sockKaoshi.listen(200)
    while int_xueshengKaoshi.get() == 1:
        try:
            conn, addr = sockKaoshi.accept()
        except:
            return
        t_Kaoshi = threading.Thread(target=thread_xueshengKaoshi, args=(conn,))
        t_Kaoshi.start()       
    sockKaoshi.close()


# 開始考試按鈕
def buttonStartKaoshiClick():
    if int_zhuce.get() == 0:
        tkinter.messagebox.showerror('很抱歉', '請聯繫作者進行軟體註冊！')
        return
    if int_xueshengKaoshi.get() == 1:
        tkinter.messagebox.showerror('很抱歉', '現在正在進行考試，不用重複設定！')
        return
    int_xueshengKaoshi.set(1)
    tkinter.messagebox.showinfo('恭喜', '設定成功，現在開始考試！')
    t_Kaoshi = threading.Thread(target=thread_xueshengKaoshiMain)
    t_Kaoshi.start()
buttonStartKaoshi = tkinter.Button(root, text='開始考試', command=buttonStartKaoshiClick)
buttonStartKaoshi.place(x=20, y=300, height=30, width=100)

# 結束考試按鈕
def buttonStopKaoshiClick():
    # 如果還沒有註冊，拒絕執行
    if int_zhuce.get() == 0:
        tkinter.messagebox.showerror('很抱歉', '請聯繫作者進行軟體註冊！')
        return
    if int_xueshengKaoshi.get() == 0:
        tkinter.messagebox.showerror('很抱歉', '現在不是考試時間！')
        return
    int_xueshengKaoshi.set(0)
    sockKaoshi.close()
    tkinter.messagebox.showinfo('恭喜', '設定成功，現在停止考試！')
buttonStopKaoshi = tkinter.Button(root, text='結束考試', command=buttonStopKaoshiClick)
buttonStopKaoshi.place(x=130, y=300, height=30, width=100)
## =====================線上考試功能程式碼結束==============================

## =====================清除考試資料功能程式碼開始==============================
# 清空所有學生考試資料
def buttonDeleteKaoshiShujuClick():
    # 如果還沒有註冊，拒絕執行
    if int_zhuce.get() == 0:
        tkinter.messagebox.showerror('很抱歉', '請聯繫作者進行軟體註冊！')
        return
    
    if tkinter.messagebox.askyesno('確定', '確定要刪除嗎？')==tkinter.YES:
        sql = "delete from kaoshi"
        Common.doSQL(sql)
        tkinter.messagebox.showinfo('恭喜', '刪除成功！')
buttonDeleteKaoshiShuju = tkinter.Button(root, text='刪除考試資料', command=buttonDeleteKaoshiShujuClick)
buttonDeleteKaoshiShuju.place(x=240, y=300, height=30, width=100)
## =====================清除考試資料功能程式碼結束==============================

## =====================產生Word版試卷功能程式碼開始==============================
def buttonGenerateShijuanClick():
    num = tkinter.simpledialog.askinteger('請輸入題目數量','題目數量')
    if not num:
        return
    conn = sqlite3.connect('database.db')
    cur = conn.cursor()
    cur.execute('select timu,daan from tiku')
    temp = cur.fetchall()
    conn.close
    temp = random.sample(temp, num)
    yesno = tkinter.messagebox.askyesno('依題型排序嗎？', '依題型排序嗎？')
    if yesno:
        # 對題目類型排序，填空題在前，是非題在後
        temp.sort(key=lambda item:item[:3], reverse=True)
    from docx import Document
    document = Document()
    document.add_paragraph('試題')
    for i, t in enumerate(temp):
        document.add_paragraph(str(i+1)+'、'+t[0])
    document.add_page_break()
    document.add_paragraph('答案')
    for i, t in enumerate(temp):
        document.add_paragraph(str(i+1)+'、'+t[1])
    document.save('試卷_答案.docx')
    tkinter.messagebox.showinfo('恭喜', '產生試卷成功！')
    ##os.startfile('試卷_答案.docx')
buttonGenerateShijuan = tkinter.Button(root, text='產生Word試卷', command=buttonGenerateShijuanClick)
buttonGenerateShijuan.place(x=20, y=340, height=30,width=100)
## =====================產生Word版試卷功能程式碼結束==============================

## =====================合併資料庫功能程式碼開始==============================
def buttonMergeDatabaseClick():
    filename = tkinter.filedialog.askopenfilename(title='請選擇SQLite資料庫檔案',
                                                  filetypes=[('sqlite Files','*.db')])
    if not filename:
        return
    if not filename.endswith('.db'):
        tkinter.messagebox.showerror('抱歉', '您必須選擇sqlite資料庫檔案！')
        return
    if os.path.abspath(filename) == os.getcwd()+'\\'+'database.db':
        tkinter.messagebox.showerror('抱歉', '您不能合併資料庫本身！')
        return

    # 開始合併資料庫
    srcConn = sqlite3.connect(filename)
    dstConn = sqlite3.connect('database.db')
    srcCur = srcConn.cursor()
    dstCur = dstConn.cursor()
    
    # 合併提問資料表
    srcCur.execute('select xuehao,shijian,defen from tiwen')
    dstCur.execute('select xuehao,shijian,defen from tiwen')
    dstRecords = dstCur.fetchall()
    for record in srcCur.fetchall():
        if record not in dstRecords:
            sqlMerge = 'insert into tiwen(xuehao,shijian,defen) values("'+\
                       record[0]+'","'+record[1]+'",'+str(record[2])+')'
            Common.doSQL(sqlMerge)

    # 合併點名資料表
    srcCur.execute('select mac,ip,xuehao,shijian from dianming')
    dstCur.execute('select mac,ip,xuehao,shijian from dianming')
    dstRecords = dstCur.fetchall()
    for record in srcCur.fetchall():
        if record not in dstRecords:
            try:
                sqlMerge = 'insert into dianming(mac,ip,xuehao,shijian) values("'+\
                           record[0]+'","'+record[1]+'","'+record[2]+'","'+record[3]+'")'
                Common.doSQL(sqlMerge)
            except:
                print(record)

    # 合併考試資訊資料表
    # 合併點名資料表
    srcCur.execute('select shijian,shifouzhengque,biaozhundaan,xuehao,xingming,timubianhao,xueshengdaan from kaoshi')
    dstCur.execute('select shijian,shifouzhengque,biaozhundaan,xuehao,xingming,timubianhao,xueshengdaan from kaoshi')
    dstRecords = dstCur.fetchall()
    for record in srcCur.fetchall():
        if record not in dstRecords:
            try:
                sqlMerge = 'insert into kaoshi(shijian,shifouzhengque,biaozhundaan,xuehao,xingming,timubianhao,xueshengdaan) values("'+\
                           record[0]+'","'+record[1]+'","'+record[2]+'","'+record[3]+\
                           '","'+record[4]+'",'+record[5]+',"',+record[6]+'")'
                Common.doSQL(sqlMerge)
            except:
                print('rrr')
    tkinter.messagebox.showinfo('恭喜', '匯入成功！')
buttonMergeDatabase = tkinter.Button(root, text='合併資料庫', command=buttonMergeDatabaseClick)
buttonMergeDatabase.place(x=130, y=340, height=30, width=100)
## =====================合併資料庫功能程式碼結束==============================

## ==========螢幕廣播程式碼開始=============
broadcasting = False
def broadcast(conn):
    global broadcasting
    while broadcasting:
        time.sleep(0.3)
        image = ImageGrab.grab()
        size = image.size
        
        imageBytes = image.tobytes()
        length = len(imageBytes)

        # 通知即將開始傳送截圖
        conn.send(b'*****')
        
        fhead = struct.pack('I32sI',
                            length,
                            str(size).encode(),
                            len(str(size).encode()))
        conn.send(fhead)

        conn.send(imageBytes)
    else:
        conn.send(b'#####')
        conn.close()

def broadcastMain():
    '''廣播螢幕截圖的主執行緒函數'''
    global sockBroadCast
    sockBroadCast = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    sockBroadCast.bind(('', 10001))
    sockBroadCast.listen(150)
    while broadcasting:
        try:
            conn, addr = sockBroadCast.accept()
        except:
            sockBroadCast.close()
            return
        threading.Thread(target=broadcast, args=(conn,)).start()
    else:
        sockBroadCast.close()
    
def onbuttonStartBroadCastClick():
    global broadcasting
    broadcasting = True
    # 啟動伺服器廣播執行緒
    threading.Thread(target=broadcastMain).start()
    
    # 通知客戶端開始接收廣播
    sock = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    sock.setsockopt(socket.SOL_SOCKET, socket.SO_BROADCAST, 1)
    IP = serverIP[:serverIP.rindex('.')]+'.255'
    sock.sendto(b'startBroadCast', (IP, 10000))
    buttonStopBroadCast['state'] = 'normal'
    buttonStartBroadCast['state'] = 'disabled'
buttonStartBroadCast = tkinter.Button(root, text='開始螢幕廣播', command=onbuttonStartBroadCastClick)
buttonStartBroadCast.place(x=20, y=380, width=100, height=30)

def onbuttonStopBroadCastClick():
    global broadcasting
    broadcasting = False
    sockBroadCast.close()
    buttonStopBroadCast['state'] = 'disabled'
    buttonStartBroadCast['state'] = 'normal'
buttonStopBroadCast = tkinter.Button(root, text='結束螢幕廣播', command=onbuttonStopBroadCastClick)
buttonStopBroadCast['state'] = 'disabled'
buttonStopBroadCast.place(x=130, y=380, width=100, height=30)

## ==========螢幕廣播程式碼結束=============



## ==========遠端關閉所有學生機器程式碼開始=======
def onbuttonShutdownClick():
    result = tkinter.messagebox.askyesno('遠端關機', '確定要關閉所有學生的機器嗎？')
    if result == tkinter.YES:
        # 通知客戶端開始接收廣播
        sock = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        sock.setsockopt(socket.SOL_SOCKET, socket.SO_BROADCAST, 1)
        IP = serverIP[:serverIP.rindex('.')]+'.255'
        sock.sendto(b'shutdown', (IP, 10000))
buttonShutdown = tkinter.Button(root, text='關閉所有學生機器', command=onbuttonShutdownClick)
buttonShutdown.place(x=240, y=340, width=100, height=30)

## ==========逞端關閉所有學生機器程式碼結束=======
root.mainloop()
